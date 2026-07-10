"""Parse hardware tracking documents from pipeline/netapp_import."""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

import pandas as pd
from docx import Document

from inventory_normalize import (
    MANUFACTURER,
    classify_event_type,
    clean_text,
    display_serial,
    format_revision,
    normalize_tool,
    parse_part_revision,
    short_serial,
    split_datetime,
)

SERIAL_TOKEN_RE = re.compile(r"\bSN[A-Z]{0,3}\d+[A-Z]?\b", re.IGNORECASE)
SERIAL_LINE_RE = re.compile(
    r"^\s*(SN[A-Z]{0,3}\d+[A-Z]?)\s*$",
    re.IGNORECASE,
)
DATE_IN_TEXT_RE = re.compile(
    r"(\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}"
    r"|\d{4}-\d{2}-\d{2}|\d{1,2}/\d{1,2}/\d{4})",
    re.IGNORECASE,
)


@dataclass
class BoardPatch:
    serial: str
    product_name: str | None = None
    board_name: str | None = None
    part_number: str | None = None
    revision: str | None = None
    tool: str | None = None
    role: str | None = None
    board_slot: str | None = None
    status: str | None = None
    comment: str | None = None
    location: str | None = None
    source_ref: str = ""


@dataclass
class BoardEvent:
    serial: str
    description: str
    event_date: str
    event_time: str | None = None
    event_type: str = "note"
    tool: str | None = None
    source_ref: str = ""


@dataclass
class ParseResult:
    path: str
    patches: list[BoardPatch] = field(default_factory=list)
    events: list[BoardEvent] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


def is_plausible_serial(value: str | None) -> bool:
    text = clean_text(value)
    if not text:
        return False
    text = text.upper().replace(" ", "")
    if text.startswith("SN"):
        text = text[2:]
    if len(text) < 3 or "." in text:
        return False
    if not re.search(r"\d", text):
        return False
    if re.match(r"^[A-Z]{1,3}\d{2,}[A-Z]?$", text):
        return True
    if re.match(r"^\d{3,}[A-Z]?$", text):
        return True
    if re.match(r"^[A-Z]\d{3,}[A-Z]?$", text):
        return True
    return False


def normalize_doc_serial(value: str | None) -> str | None:
    text = clean_text(value)
    if not text:
        return None
    text = text.upper().replace(" ", "")
    if not text.startswith("SN"):
        text = f"SN{text}"
    return text if is_plausible_serial(text) else None


def serial_core_key(value: str | None) -> str | None:
    """Normalize SN prefix and leading zeros: SND0003 and SND003 -> D3."""
    raw = clean_text(value)
    if not raw:
        return None
    text = raw.upper().replace(" ", "")
    if ":" in text:
        text = text.split(":", 1)[1]
    if text.startswith("SN"):
        text = text[2:]
    m = re.match(r"^([A-Z]*)(\d+)([A-Z]*)$", text)
    if not m:
        return text or None
    return f"{m.group(1)}{int(m.group(2))}{m.group(3)}"


def serial_match_keys(value: str | None) -> set[str]:
    text = normalize_doc_serial(value)
    if not text:
        raw = clean_text(value)
        if not raw:
            return set()
        text = raw.upper().replace(" ", "")
        if not text.startswith("SN"):
            text = f"SN{text}"
        if not is_plausible_serial(text):
            return set()
    keys = {text}
    displayed = display_serial(text)
    if displayed:
        keys.add(displayed.upper())
        keys.add(f"SN{displayed.upper()}")
    body = text[2:] if text.startswith("SN") else text
    keys.add(body)
    return {k for k in keys if k and not re.fullmatch(r"\d{1,3}", k)}


def tool_from_path(path: Path) -> str | None:
    for part in path.parts:
        m = re.search(r"tool[-_ ]?(\d+)", part, re.IGNORECASE)
        if m:
            return f"Tool{m.group(1)}"
    return None


def product_context(path: Path) -> tuple[str, str, str | None]:
    joined = "/".join(part.lower() for part in path.parts)
    if "em1_1400-10026" in joined:
        return "EM1", "LSCC", "1400-10026-09"
    if "adc1_1400-10028" in joined:
        return "ADC", "ADC", "1400-10028-03"
    if "obj1_1400-10030" in joined:
        return "OBJ", "LSCC", "1400-10030-01"
    if "bap1_1400-00064" in joined:
        return "BAP", "BAP", "1400-00064-00"
    if "bap2_1400-00064" in joined or "bap_1400-00064" in joined:
        return "BAP", "BAP2", "1400-00064-01"
    if "pwr_entry" in joined:
        return "Column", "Power Entry", "1400-20048-01"
    return "Unknown", "Unknown", None


def inventory_serial_for(product_name: str, serial: str) -> str:
    prefix_map = {
        "EM1": "EM",
        "ADC": "ADC",
        "OBJ": "Obj",
        "BAP": "",
        "Column": "",
    }
    prefix = prefix_map.get(product_name, product_name)
    sn = normalize_doc_serial(serial) or serial
    if prefix:
        return f"{prefix}:{sn}"
    return sn


def location_to_status(where: str | None) -> str | None:
    text = clean_text(where)
    if not text:
        return None
    lower = text.lower()
    if lower == "lab":
        return "Active"
    if lower.startswith("tool"):
        return "Installed"
    if "warehouse" in lower or "rack" in lower:
        return "In Stock"
    return text


def parse_event_date(text: str) -> str:
    match = DATE_IN_TEXT_RE.search(text)
    if not match:
        return pd.Timestamp.today().date().isoformat()
    event_date, _ = split_datetime(match.group(1))
    return event_date or pd.Timestamp.today().date().isoformat()


def parse_tool_assignment_xlsx(path: Path, source_ref: str) -> ParseResult:
    result = ParseResult(path=str(path))
    product_name, board_name, part_number = product_context(path)
    tool = tool_from_path(path)
    df = pd.read_excel(path, sheet_name=0, header=None)

    header_row = None
    for idx, row in df.iterrows():
        cells = [clean_text(v) for v in row.tolist()]
        if cells and str(cells[0]).lower() == "slot" and any(
            c and "role" in c.lower() for c in cells
        ):
            header_row = idx
            break

    if header_row is not None:
        headers = [clean_text(v) or "" for v in df.iloc[header_row].tolist()]
        col = {h.lower(): i for i, h in enumerate(headers)}
        role_i = next((i for h, i in col.items() if "role" in h), None)
        board_i = next((i for h, i in col.items() if h == "board"), None)
        slot_i = col.get("slot")
        new_i = next((i for h, i in col.items() if "new assignment" in h), None)
        for _, row in df.iloc[header_row + 1 :].iterrows():
            serial = normalize_doc_serial(row.iloc[board_i] if board_i is not None else None)
            if not serial:
                continue
            role = clean_text(row.iloc[role_i]) if role_i is not None else None
            slot = clean_text(row.iloc[slot_i]) if slot_i is not None else None
            new_serial = (
                normalize_doc_serial(row.iloc[new_i]) if new_i is not None else None
            )
            result.patches.append(
                BoardPatch(
                    serial=serial,
                    product_name=product_name,
                    board_name=board_name,
                    part_number=part_number,
                    tool=tool,
                    role=role,
                    board_slot=slot,
                    source_ref=source_ref,
                )
            )
            if tool:
                result.events.append(
                    BoardEvent(
                        serial=serial,
                        event_date=pd.Timestamp.today().date().isoformat(),
                        event_type="install",
                        description=f"Assigned to {tool}" + (f" ({role})" if role else ""),
                        tool=tool,
                        source_ref=f"{source_ref}:assign",
                    )
                )
            if new_serial and new_serial != serial:
                result.events.append(
                    BoardEvent(
                        serial=serial,
                        event_date=pd.Timestamp.today().date().isoformat(),
                        event_type="note",
                        description=f"Replacement assignment noted: {new_serial}",
                        tool=tool,
                        source_ref=f"{source_ref}:replacement",
                    )
                )
        return result

    for _, row in df.iterrows():
        for value in row.tolist():
            serial = normalize_doc_serial(clean_text(value))
            if serial and SERIAL_TOKEN_RE.fullmatch(serial):
                result.patches.append(
                    BoardPatch(
                        serial=serial,
                        product_name=product_name,
                        board_name=board_name,
                        part_number=part_number,
                        tool=tool,
                        source_ref=source_ref,
                    )
                )
                if tool:
                    result.events.append(
                        BoardEvent(
                            serial=serial,
                            event_date=pd.Timestamp.today().date().isoformat(),
                            event_type="install",
                            description=f"Listed on {tool} assignment sheet",
                            tool=tool,
                            source_ref=f"{source_ref}:list",
                        )
                    )
    return result


def parse_em1_repair_by_part(path: Path, source_ref: str) -> ParseResult:
    result = ParseResult(path=str(path))
    df = pd.read_excel(path, sheet_name=0, header=0)
    df.columns = [clean_text(c) or f"col{i}" for i, c in enumerate(df.columns)]
    for _, row in df.iterrows():
        serial = normalize_doc_serial(row.get("Board SN") or row.iloc[0])
        reference = clean_text(row.get("Reference") or row.iloc[2])
        part_type = clean_text(row.get("Type") or row.iloc[4])
        note = clean_text(row.get("Note") or row.iloc[-1])
        if not serial:
            continue
        parts = [p for p in (reference, part_type, note) if p]
        if not parts:
            continue
        description = "Repair by part: " + " — ".join(parts)
        result.events.append(
            BoardEvent(
                serial=serial,
                event_date=pd.Timestamp.today().date().isoformat(),
                event_type="repair",
                description=description,
                source_ref=f"{source_ref}:{serial}:{reference or 'part'}",
            )
        )
    return result


def parse_em1_board_status(path: Path, source_ref: str) -> ParseResult:
    result = ParseResult(path=str(path))
    df = pd.read_excel(path, sheet_name=0, header=None)
    header_row = None
    for idx, row in df.iterrows():
        if clean_text(row.iloc[0]) == "SN":
            header_row = idx
            break
    if header_row is None:
        result.warnings.append("Could not find SN/Where/Notes header row")
        return result

    for _, row in df.iloc[header_row + 1 :].iterrows():
        serial = normalize_doc_serial(row.iloc[0])
        where = clean_text(row.iloc[1])
        notes = clean_text(row.iloc[2])
        if not serial:
            continue
        status = location_to_status(where)
        result.patches.append(
            BoardPatch(
                serial=serial,
                product_name="EM1",
                board_name="LSCC",
                status=status,
                comment=notes,
                location=where,
                source_ref=source_ref,
            )
        )
        if where or notes:
            desc = " — ".join(p for p in (f"Location: {where}" if where else None, notes) if p)
            result.events.append(
                BoardEvent(
                    serial=serial,
                    event_date=pd.Timestamp.today().date().isoformat(),
                    event_type=classify_event_type(desc),
                    description=desc,
                    source_ref=f"{source_ref}:{serial}:status",
                )
            )
    return result


def parse_batch_numbers_xlsx(path: Path, source_ref: str) -> ParseResult:
    result = ParseResult(path=str(path))
    product_name, board_name, default_part = product_context(path)
    df = pd.read_excel(path, sheet_name=0, header=None)
    for _, row in df.iterrows():
        part_cell = clean_text(row.iloc[0])
        rev_cell = clean_text(row.iloc[1]) if len(row) > 1 else None
        serial_range = clean_text(row.iloc[2]) if len(row) > 2 else None
        if not serial_range or not SERIAL_TOKEN_RE.search(serial_range):
            continue
        part_number = part_cell if part_cell and re.match(r"\d{4}-\d{5}", part_cell) else default_part
        revision = format_revision(rev_cell) if rev_cell else None
        manufacturer = clean_text(row.iloc[3]) if len(row) > 3 else None
        note = " — ".join(
            p for p in (f"Serial range {serial_range}", manufacturer) if p
        )
        result.events.append(
            BoardEvent(
                serial=serial_range,
                event_date=pd.Timestamp.today().date().isoformat(),
                event_type="note",
                description=f"Batch record: {note}",
                source_ref=f"{source_ref}:batch:{serial_range}",
            )
        )
        for token in SERIAL_TOKEN_RE.findall(serial_range):
            serial = normalize_doc_serial(token)
            if not serial:
                continue
            result.patches.append(
                BoardPatch(
                    serial=serial,
                    product_name=product_name,
                    board_name=board_name,
                    part_number=part_number,
                    revision=revision,
                    comment=note,
                    source_ref=source_ref,
                )
            )
    return result


def parse_wire_rack_inventory(path: Path, source_ref: str) -> ParseResult:
    result = ParseResult(path=str(path))
    xl = pd.ExcelFile(path)
    for sheet in xl.sheet_names:
        df = pd.read_excel(path, sheet_name=sheet, header=0)
        df.columns = [clean_text(c) or f"col{i}" for i, c in enumerate(df.columns)]
        for _, row in df.iterrows():
            board_name = clean_text(row.get("Board Name"))
            serial_num = clean_text(row.get("Board S/N"))
            rev = clean_text(row.get("REV"))
            location = clean_text(row.get("Location"))
            note = clean_text(row.get("Note"))
            if not board_name or not serial_num:
                continue
            part_number, revision = parse_part_revision(rev) if rev else (None, None)
            serial = f"SN{serial_num}" if serial_num.isdigit() else normalize_doc_serial(serial_num)
            if not serial:
                continue
            result.patches.append(
                BoardPatch(
                    serial=serial,
                    product_name="Column",
                    board_name=board_name,
                    part_number=part_number,
                    revision=revision or format_revision(rev),
                    status=location_to_status(location),
                    comment=note,
                    location=location,
                    source_ref=f"{source_ref}:{sheet}",
                )
            )
    return result


def parse_warehouse_inventory_docx(path: Path, source_ref: str) -> ParseResult:
    result = ParseResult(path=str(path))
    doc = Document(path)
    for table in doc.tables:
        headers = [clean_text(c.text).lower() for c in table.rows[0].cells]
        if "serial" not in "".join(headers):
            continue
        serial_i = next(i for i, h in enumerate(headers) if "serial" in h)
        board_i = next((i for i, h in enumerate(headers) if "board" in h), 0)
        desc_i = next((i for i, h in enumerate(headers) if "description" in h), None)
        for row in table.rows[1:]:
            cells = [clean_text(c.text) for c in row.cells]
            if not any(cells):
                continue
            serial_raw = cells[serial_i] if serial_i < len(cells) else None
            serial = normalize_doc_serial(serial_raw.replace("SN", "").strip() if serial_raw else None)
            if not serial:
                continue
            board_code = cells[board_i] if board_i < len(cells) else None
            part_number, revision = parse_part_revision(board_code) if board_code else (None, None)
            description = cells[desc_i] if desc_i is not None and desc_i < len(cells) else None
            result.patches.append(
                BoardPatch(
                    serial=serial,
                    product_name="Column",
                    board_name="Power Entry",
                    part_number=part_number,
                    revision=revision,
                    status="In Stock",
                    comment=description,
                    source_ref=source_ref,
                )
            )
    return result


def parse_unit_repairs_docx(path: Path, source_ref: str) -> ParseResult:
    result = ParseResult(path=str(path))
    serial = normalize_doc_serial(path.stem.replace("_Repairs", ""))
    doc = Document(path)
    for table in doc.tables:
        for row in table.rows[1:]:
            cells = [clean_text(c.text) for c in row.cells]
            if len(cells) < 2:
                continue
            row_serial = normalize_doc_serial(cells[0]) or serial
            history = cells[1]
            if not row_serial or not history:
                continue
            result.events.append(
                BoardEvent(
                    serial=row_serial,
                    event_date=parse_event_date(history),
                    event_type="repair",
                    description=history,
                    source_ref=f"{source_ref}:{row_serial}",
                )
            )
    return result


def parse_paragraph_repair_log(path: Path, source_ref: str) -> ParseResult:
    result = ParseResult(path=str(path))
    product_name, board_name, part_number = product_context(path)
    doc = Document(path)
    current_serial: str | None = None
    buffer: list[str] = []

    def flush():
        nonlocal current_serial, buffer
        if not current_serial or not buffer:
            buffer = []
            return
        description = " ".join(buffer).strip()
        if description:
            result.events.append(
                BoardEvent(
                    serial=current_serial,
                    event_date=parse_event_date(description),
                    event_type="repair",
                    description=description,
                    source_ref=f"{source_ref}:{current_serial}:{len(result.events)}",
                )
            )
        buffer = []

    for para in doc.paragraphs:
        text = clean_text(para.text)
        if not text:
            continue
        if SERIAL_LINE_RE.match(text):
            flush()
            current_serial = normalize_doc_serial(text)
            continue
        if SERIAL_TOKEN_RE.fullmatch(text.replace(" ", "")):
            flush()
            current_serial = normalize_doc_serial(text)
            continue
        if current_serial:
            buffer.append(text)
        elif product_name == "BAP" and re.match(r"^SN\d+$", text, re.IGNORECASE):
            flush()
            current_serial = normalize_doc_serial(text)
    flush()

    for table in doc.tables:
        headers = [clean_text(c.text) or "" for c in table.rows[0].cells]
        lower_headers = [h.lower() for h in headers]
        if lower_headers and "serial number" in lower_headers[0]:
            for row in table.rows[1:]:
                cells = [clean_text(c.text) for c in row.cells]
                if not cells or not cells[0]:
                    continue
                serial = normalize_doc_serial(cells[0])
                note = " — ".join(c for c in cells[1:] if c)
                if serial and note:
                    result.events.append(
                        BoardEvent(
                            serial=serial,
                            event_date=parse_event_date(note),
                            event_type="repair",
                            description=note,
                            source_ref=f"{source_ref}:{serial}:table",
                        )
                    )
            continue
        if headers and headers[0].lower() == "serial number":
            for row in table.rows[1:]:
                cells = [clean_text(c.text) for c in row.cells]
                if len(cells) < 4:
                    continue
                serial = normalize_doc_serial(cells[0])
                note = cells[3]
                if serial and note:
                    result.patches.append(
                        BoardPatch(
                            serial=serial,
                            product_name=product_name,
                            board_name=board_name,
                            part_number=part_number,
                            revision=format_revision(cells[1]),
                            tool=normalize_tool(cells[2]),
                            source_ref=source_ref,
                        )
                    )
                    result.events.append(
                        BoardEvent(
                            serial=serial,
                            event_date=parse_event_date(note),
                            event_type="repair",
                            description=note,
                            tool=normalize_tool(cells[2]),
                            source_ref=f"{source_ref}:{serial}:log",
                        )
                    )
    return result


def parse_em1_batch_xlsx(path: Path, source_ref: str) -> ParseResult:
    """Extract serial numbers from batch test workbook column headers."""
    result = ParseResult(path=str(path))
    df = pd.read_excel(path, sheet_name=0, header=None, nrows=3)
    seen: set[str] = set()
    for value in df.iloc[0].tolist():
        serial = normalize_doc_serial(clean_text(value))
        if not serial or not SERIAL_TOKEN_RE.fullmatch(serial):
            continue
        if serial in seen:
            continue
        seen.add(serial)
        result.patches.append(
            BoardPatch(
                serial=serial,
                product_name="EM1",
                board_name="LSCC",
                part_number="1400-10026-09",
                source_ref=source_ref,
            )
        )
    return result


def classify_file(path: Path) -> str:
    name = path.name.lower()
    if name.endswith("_repairs.docx"):
        return "unit_repairs"
    if "repair" in name and name.endswith(".docx"):
        return "repair_log"
    if "warehouse inventory" in name:
        return "warehouse_inventory"
    if "wire rack cage_board inventory" in name:
        return "wire_rack_inventory"
    if "batchnumbers" in name:
        return "batch_numbers"
    if name.startswith("em1_batch_") and name.endswith(".xlsx"):
        return "em1_batch"
    if name == "em1 repair by part.xlsx":
        return "em1_repair_by_part"
    if name == "em1_board_status.xlsx":
        return "em1_board_status"
    if "assignment" in name and name.endswith(".xlsx"):
        return "tool_assignment"
    return "unknown"


def parse_file(path: Path) -> ParseResult:
    source_ref = str(path.as_posix())
    kind = classify_file(path)
    if kind == "tool_assignment":
        return parse_tool_assignment_xlsx(path, source_ref)
    if kind == "em1_repair_by_part":
        return parse_em1_repair_by_part(path, source_ref)
    if kind == "em1_board_status":
        return parse_em1_board_status(path, source_ref)
    if kind == "batch_numbers":
        return parse_batch_numbers_xlsx(path, source_ref)
    if kind == "wire_rack_inventory":
        return parse_wire_rack_inventory(path, source_ref)
    if kind == "warehouse_inventory":
        return parse_warehouse_inventory_docx(path, source_ref)
    if kind == "unit_repairs":
        return parse_unit_repairs_docx(path, source_ref)
    if kind == "repair_log":
        return parse_paragraph_repair_log(path, source_ref)
    if kind == "em1_batch":
        return parse_em1_batch_xlsx(path, source_ref)
    return ParseResult(path=str(path), warnings=[f"Unhandled file type: {path.name}"])


def iter_import_files(root: Path) -> list[Path]:
    files: list[Path] = []
    for path in sorted(root.rglob("*")):
        if not path.is_file():
            continue
        if path.name.startswith("~$") or path.name == "_copy_manifest.json":
            continue
        if path.suffix.lower() not in {".xlsx", ".xls", ".docx"}:
            continue
        files.append(path)
    return files


def parse_import_tree(root: Path) -> list[ParseResult]:
    return [parse_file(path) for path in iter_import_files(root)]
